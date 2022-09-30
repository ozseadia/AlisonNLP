# -*- coding: utf-8 -*-
"""
Created on Fri Aug  5 17:31:42 2022

@author: OzSea
"""
import time
import numpy as np 
import requests
import json
import docx
import re
import subprocess
import threading
from transformers import AutoTokenizer, AutoModel, pipeline
#import Results2Doc as RD
import Results2excel as RD
import openpyxl

class FEYAP ():
    def __init__(self):
        self.p=subprocess.Popen(['C://Users//OzSea.LAPTOP-LLBIIFTU//yapproj//src//yap//yap','api'], shell=False, stdout=subprocess.PIPE)
        time.sleep(30)
        #subprocess.run(['C://Users//OzSea//yapproj//src//yap//yap','api'])
        self.localhost_yap = "http://localhost:8000/yap/heb/joint"
        self.headers = {'content-type': 'application/json'}
        # while (self.p.poll()):
        #     time.sleep(1)
        self.sentiment_analysis = pipeline(
	    "sentiment-analysis",
	    model="avichr/heBERT_sentiment_analysis",
	    tokenizer="avichr/heBERT_sentiment_analysis",
	    return_all_scores = True)
            
    def Text2sentenc(self,text):
        data = json.dumps({'text': "{}  ".format(text)})  # input string ends with two space characters
        response = requests.get(url=self.localhost_yap, data=data, headers=self.headers)
        json_response = response.json()
        return(json_response)
    
    def RepetWord1(self,TT):
        
        for i in range(len(TT)):
            TT[i]=TT[i].split('\t')
        n=0
        for i in range(len(TT)-1):
            
            if len(TT[i+1])>=3 and len(TT[i])>=3:
                #print(i)
                try:
                    n+=(TT[i][2]==TT[i+1][2]) 
                except:
                    print(TT[i+1])
        return(n) 
    
    def RepetWord3(self,text):
        Temp=re.sub(r'[^\w]', ' ', text) # remove all symbols from string
        Temp = list(filter(None, Temp.split(' ')))
        Words_List=list(set(Temp))
        words_dict = dict.fromkeys(Words_List,0)
        for key in words_dict:
            words_dict[key]=(re.sub(r'[^\w]', ' ', text)).count(' '+key+' ')
            
        return(words_dict)
    
    def CloseYap(self):
        self.p.terminate()
    
    def Sentiment(self,text_raw,Nwords):
        ListSentence=text_raw.split('.')
        subtext=list()
        for text in ListSentence:
            if len(text.split(' '))<=Nwords:
                subtext.append(text)
            else:
                
                R=int((len(text.split(' '))-len(text.split(' '))%Nwords)/Nwords)
                for i in range(R+1):
                    Sen=''
                    if i<R:
                        for j in range(Nwords*i,Nwords*(i+1)):
                            Sen+=' '+text.split(' ')[j]
                        subtext.append(Sen[1:])
                    else:
                        for j in range(Nwords*i,len(text.split(' '))):
                            Sen+=' '+text.split(' ')[j]
                        subtext.append(Sen[1:]) 
        S=self.sentiment_analysis(subtext)   
        Score=dict({'negative':0,'positive':0,'neutral':0})
        for e in S:
            for i in range(len(e)):
                Score[list(e[i].values())[0]]+=list(e[i].values())[1]/len(S)
        print(subtext)        
        return(Score)
        
    def ExtractWord(self,text):
        pass
    
    def FE_Yap(self,Text_path="101 - Atlas.docx"):

        doc = docx.Document(Text_path)
        fullText = str()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color.rgb:
                    print(run.font.color.rgb)
                else:
                    #fullText.append(run.text)
                    fullText+=run.text
        
        #result = [p.text for p in doc.paragraphs] 
        result=[fullText]
        result[0]=re.sub("\(.*?\)","",result[0])
        result[0]=re.sub("\[.*?\]","",result[0])
        result[0]=re.sub("[\[\]<>()\\/\#]","",result[0])
        RepeatWords3=self.RepetWord3(result[0])
        Sentiment_Score=self.Sentiment(result[0])
        NumberOfwords=result[0].count(' ')#total number of words in the text
        TEXT=result[0].split('.')
        presentCount=0
        pastCount=0
        futureCount=0
        Gufim=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'They':0,'He':0,'She':0})
        Conjugation=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'They':0,'He':0,'She':0})
        GufimWords=list()
        ConjugationWords=list()
        NWqm=0
        RepeatWords=0
        #******************************************************************
        #What percentage of words are enclosed in quotation marks (markers of dialogue)
        Quotation_marks= [_.start() for _ in re.finditer('"', result[0])]
        Nqm=0
        print(Quotation_marks)
        if Quotation_marks:
            for i in range(0,len(Quotation_marks),2):
                t=result[0][Quotation_marks[i]:Quotation_marks[i+1]+1]
                Nqm+=len(t.split(' '))
            NWqm+=Nqm/NumberOfwords
        
        #******************************************************************
        TABLE=list()
        for text in TEXT:
            data = json.dumps({'text': "{}  ".format(text)})  # input string ends with two space characters
            response = requests.get(url=self.localhost_yap, data=data, headers=self.headers)
            json_response = response.json()
            RepeatWords+=self.RepetWord1(json_response['dep_tree'].split('\n'))
            Table=json_response['dep_tree'].split('\n')
            #print(Table)
            TABLE.append(Table)
            
            
            #*********************************************************************
            #Feature1: What percentage of total verbs appear in past, present, future
            
            presentCount+= json_response['md_lattice'].count('tense=BEINONI')
            pastCount += json_response['md_lattice'].count('tense=PAST')  
            futureCount += json_response['md_lattice'].count('tense=FUTURE')
            #temp=presentCount+pastCount+futureCount
            temp=1
            if temp>0:
                Tense=dict({'Past':pastCount/temp,'Present':presentCount/temp,
                        'Future':futureCount/temp})
            

            
            for i in range(len(Table)):
                #print(type(Table[i]))
                if not(re.findall(r'suf_per', Table[i], re.IGNORECASE)):
                    if not (re.findall(r'S_PRN', Table[i], re.IGNORECASE)):
                        match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                        Gufim['Me']+=len(match_list)==2
                        #if len(match_list)==2:GufimWords.append(Table[i])
                        #if len(match_list)==2:
                        #     print(str(Table[i].split('\t')))
                        match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                        Gufim['YouMs']+=len(match_list)==3
                        #if len(match_list)==3:GufimWords.append(Table[i])
                        match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                        Gufim['YouFs']+=len(match_list)==3
                        #if len(match_list)==3:GufimWords.append(Table[i])
                        match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                        Gufim['Youp']+=len(match_list)==2
                        #if len(match_list)==2:GufimWords.append(Table[i])
                        match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                        Gufim['We']+=len(match_list)==2
                        #if len(match_list)==2:GufimWords.append(Table[i])
                        match_list = re.findall(r'num=P|per=3', Table[i], re.IGNORECASE)
                        Gufim['They']+=len(match_list)==2
                        #if len(match_list)==2:GufimWords.append(Table[i])
                        #print(match_list)
                        #if len(match_list)==2:
                        #    print(match_list)
                        #    print(str(Table[i].split('\t')))
                        match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                        Gufim['He']+=len(match_list)==3
                        #if len(match_list)==3:GufimWords.append(Table[i])
                        match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                        Gufim['She']+=len(match_list)==3
                        match_list = re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE)
                        if (re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE) and 
                            re.findall(r'num', Table[i], re.IGNORECASE)):GufimWords.append(Table[i])
                        #if len(match_list)==1:GufimWords.append(Table[i])
                    if (re.findall(r'S_PRN', Table[i], re.IGNORECASE)):
                        match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                        Conjugation['Me']+=len(match_list)==2
                        match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                        Conjugation['YouMs']+=len(match_list)==3
                        match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                        Conjugation['YouFs']+=len(match_list)==3
                        match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                        Conjugation['Youp']+=len(match_list)==2
                        match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                        Conjugation['We']+=len(match_list)==2
                        match_list = re.findall(r'per=3|num=P', Table[i], re.IGNORECASE)
                        Conjugation['They']+=len(match_list)==2
                        
                        match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                        Conjugation['He']+=len(match_list)==3
                        match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                        Conjugation['She']+=len(match_list)==3
                        ConjugationWords.append(Table[i])
                    #match_list = re.findall(r'S_PRN', Table[i], re.IGNORECASE)
                    #if len(match_list)>=1:ConjugationWords.append(Table[i])

        temp=1
        #temp=sum(Gufim.values())
        for id in Gufim:
            Gufim[id]=Gufim[id]/temp
        #temp=sum(Conjugation.values())
        temp=1
        for id in Conjugation:
            Conjugation[id]=Conjugation[id]/temp
        Temp=str()
        for i in range(len(GufimWords)):
            Temp+=' '+GufimWords[i].split('\t')[1]
            Temp+=' '
        GW=self.RepetWord3(Temp)    
        Temp=str()
        for i in range(len(ConjugationWords)):
            Temp+=' '+ConjugationWords[i].split('\t')[1]
            Temp+=' '
        CW=self.RepetWord3(Temp)
        
        
        
        RD.Results(Text_path[0:-5],Sentiment_Score,RepeatWords3,RepeatWords,Tense,Gufim,Conjugation,NWqm,GW,CW)
        return (TABLE,ConjugationWords,GufimWords,Sentiment_Score,RepeatWords3,RepeatWords,Tense,Gufim,Conjugation,NWqm)
            
    def FE_Yap1(self,Text_path="101 - Atlas.docx"):
        
        doc = docx.Document(Text_path)
        fullText = str()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color.rgb:
                    print(run.font.color.rgb)
                else:
                    #fullText.append(run.text)
                    fullText+=run.text
        
        #result = [p.text for p in doc.paragraphs]
        Nwords=1000
        result=[fullText]
        result[0]=re.sub("\(.*?\)","",result[0])
        result[0]=re.sub("\[.*?\]","",result[0])
        result[0]=re.sub("[\[\]<>()\\/\#]","",result[0])
        RepeatWords3=self.RepetWord3(result[0])
        Sentiment_Score=self.Sentiment(result[0],5)
        NumberOfwords=result[0].count(' ')#total number of words in the text
        TEXT=result[0].split('.')
        presentCount=0
        pastCount=0
        futureCount=0
        Test=''
        GufimPast=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        GufimPastWords=dict({'Me':[],'YouMs':[],'YouFs':[],'Youp':[],'We':[],'He':[],'She':[],'They':[]})
        GufimPresent=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        GufimPresentWords=dict({'Me':[],'YouMs':[],'YouFs':[],'Youp':[],'We':[],'He':[],'She':[],'They':[]})
        GufimFuture=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        GufimFutureWords=dict({'Me':[],'YouMs':[],'YouFs':[],'Youp':[],'We':[],'He':[],'She':[],'They':[]})
        GufimPrepositions=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        GufimPrepositionsWords=dict({'Me':[],'YouMs':[],'YouFs':[],'Youp':[],'We':[],'He':[],'She':[],'They':[]})
        Gufim=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        COP=dict({'Me':0,'YouMs':0,'YouFs':0,'Youp':0,'We':0,'He':0,'She':0,'They':0})
        GufimTenseWords=list()
        #GufimPrepositionsWords=list()
        GufimWords=list()
        ConjugationWords=list()
        COPWords=list()
        NWqm=0
        RepeatWords=0
        #******************************************************************
        #What percentage of words are enclosed in quotation marks (markers of dialogue)
        Quotation_marks= [_.start() for _ in re.finditer('"', result[0])]
        Nqm=0
        print(Quotation_marks)
        if Quotation_marks:
            for i in range(0,len(Quotation_marks),2):
                t=result[0][Quotation_marks[i]:Quotation_marks[i+1]+1]
                Nqm+=len(t.split(' '))
            NWqm+=Nqm/NumberOfwords
        
        #******************************************************************
        TABLE=list()
        
        for text in TEXT:
            if len(text.split(' '))<=Nwords:
                data = json.dumps({'text': "{}  ".format(text)})  # input string ends with two space characters
                response = requests.get(url=self.localhost_yap, data=data, headers=self.headers)
                json_response = response.json()
                RepeatWords+=self.RepetWord1(json_response['dep_tree'].split('\n'))
                Table=json_response['dep_tree'].split('\n')
                #Table=json_response['md_lattice'].split('\n')
                #print(Table)
                TABLE.append(Table)
            else:
                Ta=list()
                R=int((len(text.split(' '))-len(text.split(' '))%Nwords)/Nwords)
                for i in range(R+1):
                    S=''
                    if i<R:
                        for j in range(Nwords*i,Nwords*(i+1)):
                            S+=' '+text.split(' ')[j]
                        subtext=S[1:]
                    else:
                        for j in range(Nwords*i,len(text.split(' '))):
                            S+=' '+text.split(' ')[j]
                        subtext=S[1:]
                    data = json.dumps({'text': "{}  ".format(subtext)})  # input string ends with two space characters
                    response = requests.get(url=self.localhost_yap, data=data, headers=self.headers)
                    json_response = response.json()
                    RepeatWords+=self.RepetWord1(json_response['dep_tree'].split('\n'))
                    Ta.append(json_response['dep_tree'].split('\n'))
                #Table=[]    
                Table = [item for sublist in Ta for item in sublist] 
                TABLE.append(Table)
            
            
            #*********************************************************************
            #Feature1: What percentage of total verbs appear in past, present, future
            
            presentCount+= json_response['md_lattice'].count('tense=BEINONI')
            pastCount += json_response['md_lattice'].count('tense=PAST')  
            futureCount += json_response['md_lattice'].count('tense=FUTURE')
            #temp=presentCount+pastCount+futureCount
            temp=1
            if temp>0:
                Tense=dict({'Past':pastCount/temp,'Present':presentCount/temp,
                        'Future':futureCount/temp})
            

            
            for i in range(len(Table)):
                #print(type(Table[i]))
                #print(i)
                #Table[i]=list(filter(None,Table[i]))
                #print(Table)
                if (re.findall(r'tense=PAST', Table[i], re.IGNORECASE)):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    GufimPast['Me']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['Me'].append(str(Table[i].split('\t')))
                    #if len(match_list)==2: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPast['YouMs']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['YouMs'].append(str(Table[i].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPast['YouFs']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['YouFs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    GufimPast['Youp']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['Youp'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    GufimPast['We']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['We'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=3', Table[i], re.IGNORECASE)
                    GufimPast['They']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['They'].append(str(Table[i].split('\t')))
                    #if len(match_list)==2: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPast['He']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['He'].append(str(Table[i].split('\t'))) 
                    #if len(match_list)==3: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPast['She']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['She'].append(str(Table[i].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE)
                    GufimTenseWords.append(Table[i])
                if (re.findall(r'COP', Table[i], re.IGNORECASE) and not (re.findall('אני|את|אתה|אתם|אתן|אנחנו|הם|הן|הוא|היא',
                                                                                    Table[i], re.IGNORECASE))):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    COP['Me']+=len(match_list)==2
                    GufimPast['Me']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['Me'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    COP['YouMs']+=len(match_list)==3
                    GufimPast['YouMs']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['YouMs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    COP['YouFs']+=len(match_list)==3
                    GufimPast['YouFs']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['YouFs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    COP['Youp']+=len(match_list)==2
                    GufimPast['Youp']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['Youp'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    COP['We']+=len(match_list)==2
                    GufimPast['We']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['We'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=3', Table[i], re.IGNORECASE)
                    COP['They']+=len(match_list)==2
                    GufimPast['They']+=len(match_list)==2
                    if len(match_list)==2: GufimPastWords['They'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    COP['He']+=len(match_list)==3
                    GufimPast['He']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['He'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    COP['She']+=len(match_list)==3
                    GufimPast['She']+=len(match_list)==3
                    if len(match_list)==3: GufimPastWords['She'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE)
                    COPWords.append(Table[i])
                
                if (re.findall(r'tense=BEINONI', Table[i], re.IGNORECASE)):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    GufimPresent['Me']+=len(match_list)==2
                    if len(match_list)==2: GufimPresentWords['Me'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPresent['YouMs']+=len(match_list)==3
                    if len(match_list)==3: GufimPresentWords['YouMs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPresent['YouFs']+=len(match_list)==3
                    if len(match_list)==3: GufimPresentWords['YouFs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    GufimPresent['Youp']+=len(match_list)==2
                    if len(match_list)==2: GufimPresentWords['YouP'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    GufimPresent['We']+=len(match_list)==2
                    if len(match_list)==2: GufimPresentWords['We'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=3', Table[i], re.IGNORECASE)
                    GufimPresent['They']+=len(match_list)==2
                    if len(match_list)==2: GufimPresentWords['They'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPresent['He']+=len(match_list)==3
                    if len(match_list)==3: GufimPresentWords['He'].append(str(Table[i].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPresent['She']+=len(match_list)==3
                    if len(match_list)==3: GufimPresentWords['She'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE)
                    GufimTenseWords.append(Table[i])    
                if (re.findall(r'tense=FUTURE', Table[i], re.IGNORECASE)):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    GufimFuture['Me']+=len(match_list)==2
                    if len(match_list)==2: GufimFutureWords['Me'].append(str(Table[i].split('\t')))
                    #if len(match_list)==2: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimFuture['YouMs']+=len(match_list)==3
                    if len(match_list)==3: GufimFutureWords['YouMs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimFuture['YouFs']+=len(match_list)==3
                    if len(match_list)==3: GufimFutureWords['YouFs'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    GufimFuture['Youp']+=len(match_list)==2
                    if len(match_list)==2: GufimFutureWords['Youp'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    GufimFuture['We']+=len(match_list)==2
                    if len(match_list)==2: GufimFutureWords['We'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'num=P|per=3', Table[i], re.IGNORECASE)
                    GufimFuture['They']+=len(match_list)==2
                    if len(match_list)==2: GufimFutureWords['They'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimFuture['He']+=len(match_list)==3
                    if len(match_list)==3: GufimFutureWords['He'].append(str(Table[i].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimFuture['She']+=len(match_list)==3
                    if len(match_list)==3: GufimFutureWords['She'].append(str(Table[i].split('\t')))
                    match_list = re.findall(r'per=1|per=2|per=3', Table[i], re.IGNORECASE)
                    GufimTenseWords.append(Table[i])     
                    
    
                #if not re.findall('VB', Table[i], re.IGNORECASE):    
                if (re.findall('S_PRN', Table[i], re.IGNORECASE)):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    GufimPrepositions['Me']+=len(match_list)==2
                    if len(match_list)==2: GufimPrepositionsWords['Me'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPrepositions['YouMs']+=len(match_list)==3
                    if len(match_list)==3: GufimPrepositionsWords['YouMs'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    GufimPrepositions['YouFs']+=len(match_list)==3
                    if len(match_list)==3: GufimPrepositionsWords['YouFs'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    GufimPrepositions['Youp']+=len(match_list)==2
                    if len(match_list)==2: GufimPrepositionsWords['Youp'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    GufimPrepositions['We']+=len(match_list)==2
                    if len(match_list)==2: GufimPrepositionsWords['We'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'per=3|num=P', Table[i], re.IGNORECASE)
                    GufimPrepositions['They']+=len(match_list)==2
                    if len(match_list)==2: GufimPrepositionsWords['They'].append(str(Table[i-1].split('\t')))
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPrepositions['He']+=len(match_list)==3
                    if len(match_list)==3: GufimPrepositionsWords['He'].append(str(Table[i-1].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]+''+Table[i-1]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    GufimPrepositions['She']+=len(match_list)==3
                    if len(match_list)==3: GufimPrepositionsWords['She'].append(str(Table[i-1].split('\t')))
                    #if len(match_list)==3: Test+=' '+Table[i]+''+Table[i-1]#.split('\t')[1]
                    #GufimPrepositionsWords.append(Table[i-1])

                if ((re.findall('PRP', Table[i], re.IGNORECASE)) and 
                (re.findall('אני|את|אתה|אתם|אתן|אנחנו|הם|הן|הוא|היא', Table[i], re.IGNORECASE))):
                    match_list = re.findall(r'num=S|per=1', Table[i], re.IGNORECASE)
                    Gufim['Me']+=len(match_list)==2
                    match_list = re.findall(r'gen=M|num=S|per=2', Table[i], re.IGNORECASE)
                    Gufim['YouMs']+=len(match_list)==3
                    match_list = re.findall(r'gen=F|num=S|per=2', Table[i], re.IGNORECASE)
                    Gufim['YouFs']+=len(match_list)==3
                    match_list = re.findall(r'num=P|per=2', Table[i], re.IGNORECASE)
                    Gufim['Youp']+=len(match_list)==2
                    match_list = re.findall(r'num=P|per=1', Table[i], re.IGNORECASE)
                    Gufim['We']+=len(match_list)==2
                    match_list = re.findall(r'per=3|num=P', Table[i], re.IGNORECASE)
                    Gufim['They']+=len(match_list)==2
                    match_list = re.findall(r'gen=M|num=S|per=3', Table[i], re.IGNORECASE)
                    Gufim['He']+=len(match_list)==3
                    #if len(match_list)==3: Test+=' '+Table[i]+''+Table[i-1]#.split('\t')[1]
                    match_list = re.findall(r'gen=F|num=S|per=3', Table[i], re.IGNORECASE)
                    Gufim['She']+=len(match_list)==3
                    #if len(match_list)==3: Test+=' '+Table[i]+''+Table[i-1]#.split('\t')[1]
                    GufimWords.append(Table[i])


        Temp=str()
        for i in range(len(GufimTenseWords)):
            Temp+=' '+GufimTenseWords[i].split('\t')[1]
            Temp+=' '
        GT=self.RepetWord3(Temp) 
        
        # Temp=str()
        # for i in range(len(GufimPrepositionsWords)):
        #     Temp+=' '+GufimPrepositionsWords[i].split('\t')[1]
        #     Temp+=' '
        # GP=self.RepetWord3(Temp)  
            
        # Temp=str()
        # for i in range(len(GufimWords)):
        #     Temp+=' '+GufimWords[i].split('\t')[1]
        #     Temp+=' '
        # GW=self.RepetWord3(Temp)
        
        Temp=str()
        for i in range(len(COPWords)):
            Temp+=' '+COPWords[i].split('\t')[1]
            Temp+=' '
        CO=self.RepetWord3(Temp)
        
        TotalTense=sum(GufimPast.values())+sum(GufimPresent.values())+sum(GufimFuture.values())
        my_wb = openpyxl.Workbook()
        my_wb=RD.ResulysInitial(my_wb,R=1)
        my_wb,R0=RD.Results(my_wb,1,GufimPastWords,GufimPast,TotalTense,'Past')
        my_wb,R0=RD.Results(my_wb,R0,GufimPresentWords,GufimPresent,TotalTense,'Present')
        my_wb,R0=RD.Results(my_wb,R0,GufimFutureWords,GufimFuture,TotalTense,'Future')
        my_wb,R0=RD.Results(my_wb,R0,GufimPrepositionsWords,GufimPrepositions,sum(GufimPrepositions.values()),'Preposition')
        my_wb,R0=RD.Results1(my_wb,R0,Tense)
        my_wb,R0=RD.Results2(my_wb,R0,Sentiment_Score ,'Sentiment')
        my_wb.save(Text_path[0:-5]+" Results.xlsx")
    #RD.Results(Text_path[0:-5],Sentiment_Score,RepeatWords3,RepeatWords,Tense,Gufim,Conjugation,NWqm,GW,CW)
        return (Test,CO,COP,GufimPastWords,TABLE,Tense,GT,Gufim,GufimPrepositions,GufimPast,GufimPresent,GufimFuture)
 
